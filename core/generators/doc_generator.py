"""
DOC Generator - Generate DOC documents from processed data
"""
import io
from typing import Dict, Any
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from core.generators.base_generator import BaseGenerator

class DOCGenerator(BaseGenerator):
    """Generates DOC documents from processed Excel data"""
    
    def __init__(self, data: Dict[str, Any]):
        super().__init__(data)
    
    def generate_doc_documents(self) -> Dict[str, bytes]:
        """
        Generate all required documents in DOC format
        
        Returns:
            Dictionary containing all generated documents in DOC format (bytes)
        """
        doc_documents = {}
        
        # Generate DOC versions of all documents
        doc_documents['First Page Summary.docx'] = self._generate_doc_first_page()
        doc_documents['Deviation Statement.docx'] = self._generate_doc_deviation_statement()
        doc_documents['BILL SCRUTINY SHEET.docx'] = self._generate_doc_note_sheet()
        
        # Only generate Extra Items document if there are extra items
        if self._has_extra_items():
            doc_documents['Extra Items Statement.docx'] = self._generate_doc_extra_items()
        
        doc_documents['Certificate II.docx'] = self._generate_doc_certificate_ii()
        doc_documents['Certificate III.docx'] = self._generate_doc_certificate_iii()
        
        return doc_documents
    
    def _generate_doc_first_page(self) -> bytes:
        """Generate First Page Summary document in DOC format"""
        doc = Document()
        
        # Add title
        title = doc.add_heading('First Page Summary', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        current_date = datetime.now().strftime('%d/%m/%Y')
        date_para = doc.add_paragraph(f'Date: {current_date}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add project information
        doc.add_heading('Project Information', level=1)
        
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        
        # Populate project info
        table.cell(0, 0).text = 'Project Name:'
        table.cell(0, 1).text = str(self.title_data.get('Project Name', 'N/A'))
        table.cell(1, 0).text = 'Contract No:'
        table.cell(1, 1).text = str(self.title_data.get('Contract No', 'N/A'))
        table.cell(2, 0).text = 'Work Order No:'
        table.cell(2, 1).text = str(self.title_data.get('Work Order No', 'N/A'))
        
        # Add work items summary
        doc.add_heading('Work Items Summary', level=1)
        
        # Create work items table
        if not self.work_order_data.empty:
            # Add header row
            table = doc.add_table(rows=1, cols=9)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Unit'
            hdr_cells[1].text = 'Quantity executed (or supplied) since last certificate'
            hdr_cells[2].text = 'Quantity executed (or supplied) upto date as per MB'
            hdr_cells[3].text = 'S. No.'
            hdr_cells[4].text = 'Item of Work supplies'
            hdr_cells[5].text = 'Rate'
            hdr_cells[6].text = 'Upto date Amount'
            hdr_cells[7].text = 'Amount Since previous bill'
            hdr_cells[8].text = 'Remarks'
            
            # Add data rows
            for index, row in self.work_order_data.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(row.get('Unit', ''))
                row_cells[1].text = str(row.get('Quantity Since', ''))
                row_cells[2].text = str(row.get('Quantity Upto', ''))
                row_cells[3].text = str(row.get('Item No.', ''))
                row_cells[4].text = str(row.get('Description', ''))
                row_cells[5].text = str(row.get('Rate', ''))
                row_cells[6].text = str(row.get('Amount', ''))
                row_cells[7].text = str(row.get('Amount', ''))
                row_cells[8].text = str(row.get('Remark', ''))
        
        # Add totals section
        doc.add_heading('Totals', level=1)
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()
    
    def _generate_doc_deviation_statement(self) -> bytes:
        """Generate Deviation Statement document in DOC format"""
        doc = Document()
        
        # Add title
        title = doc.add_heading('Deviation Statement', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        current_date = datetime.now().strftime('%d/%m/%Y')
        date_para = doc.add_paragraph(f'Date: {current_date}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add content
        doc.add_paragraph('This is a deviation statement document.')
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()
    
    def _generate_doc_note_sheet(self) -> bytes:
        """Generate Bill Scrutiny Sheet document in DOC format"""
        doc = Document()
        
        # Add title
        title = doc.add_heading('Bill Scrutiny Sheet', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        current_date = datetime.now().strftime('%d/%m/%Y')
        date_para = doc.add_paragraph(f'Date: {current_date}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add content
        doc.add_paragraph('This is a final bill scrutiny sheet document.')
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()
    
    def _generate_doc_extra_items(self) -> bytes:
        """Generate Extra Items Statement document in DOC format"""
        doc = Document()
        
        # Add title
        title = doc.add_heading('Extra Items Statement', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        current_date = datetime.now().strftime('%d/%m/%Y')
        date_para = doc.add_paragraph(f'Date: {current_date}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add content
        doc.add_paragraph('This is an extra items statement document.')
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()
    
    def _generate_doc_certificate_ii(self) -> bytes:
        """Generate Certificate II document in DOC format"""
        doc = Document()
        
        # Add title
        title = doc.add_heading('Certificate II - Work Completion Certificate', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        current_date = datetime.now().strftime('%d/%m/%Y')
        date_para = doc.add_paragraph(f'Date: {current_date}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add content
        doc.add_paragraph('This is a work completion certificate document.')
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()
    
    def _generate_doc_certificate_iii(self) -> bytes:
        """Generate Certificate III document in DOC format"""
        doc = Document()
        
        # Add title
        title = doc.add_heading('Certificate III - Payment Certification', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        current_date = datetime.now().strftime('%d/%m/%Y')
        date_para = doc.add_paragraph(f'Date: {current_date}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add content
        doc.add_paragraph('This is a payment certification document.')
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()