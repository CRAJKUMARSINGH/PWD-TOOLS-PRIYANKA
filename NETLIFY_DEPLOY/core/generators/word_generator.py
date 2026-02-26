"""
Word Document Generator - Generate .docx files from HTML templates
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import re
from typing import Dict, Any
from pathlib import Path


class WordGenerator:
    """Generate Word documents from HTML content"""
    
    def __init__(self):
        """Initialize Word generator"""
        pass
    
    def html_to_docx(self, html_content: str, doc_name: str) -> bytes:
        """
        Convert HTML to Word document
        
        Args:
            html_content: HTML content
            doc_name: Document name
            
        Returns:
            Word document as bytes
        """
        # Create new document
        doc = Document()
        
        # Set margins (10mm = 0.39 inches)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.39)
            section.bottom_margin = Inches(0.39)
            section.left_margin = Inches(0.39)
            section.right_margin = Inches(0.39)
        
        # Parse HTML
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Add title
        title = doc.add_heading(doc_name, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Find main content table
        tables = soup.find_all('table')
        
        for html_table in tables:
            # Extract table data
            rows = html_table.find_all('tr')
            if not rows:
                continue
            
            # Create Word table
            max_cols = max(len(row.find_all(['th', 'td'])) for row in rows)
            word_table = doc.add_table(rows=len(rows), cols=max_cols)
            word_table.style = 'Table Grid'
            
            # Fill table
            for row_idx, html_row in enumerate(rows):
                cells = html_row.find_all(['th', 'td'])
                for col_idx, cell in enumerate(cells):
                    if col_idx < max_cols:
                        word_cell = word_table.rows[row_idx].cells[col_idx]
                        
                        # Get text content
                        text = cell.get_text(strip=True)
                        word_cell.text = text
                        
                        # Format header cells
                        if cell.name == 'th':
                            for paragraph in word_cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
                                    run.font.size = Pt(9)
                        else:
                            for paragraph in word_cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(8)
            
            # Add spacing after table
            doc.add_paragraph()
        
        # Save to bytes
        from io import BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def generate_all_docx(self, html_documents: Dict[str, str]) -> Dict[str, bytes]:
        """
        Generate Word documents for all HTML documents
        
        Args:
            html_documents: Dictionary of document names to HTML content
            
        Returns:
            Dictionary of document names to Word document bytes
        """
        docx_documents = {}
        
        for doc_name, html_content in html_documents.items():
            try:
                docx_bytes = self.html_to_docx(html_content, doc_name)
                docx_documents[doc_name] = docx_bytes
            except Exception as e:
                print(f"Error generating Word document for {doc_name}: {e}")
        
        return docx_documents
